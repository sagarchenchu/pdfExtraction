"""
PDF Table Extractor - Standalone GUI Application
Extracts table data and text from PDF pages and exports to Excel, Word, or TXT.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import concurrent.futures
import os
import queue

# ---------------------------------------------------------------------------
# Optional heavy imports – caught so the app still opens even if a package is
# missing at runtime, and we can show a friendly error to the user.
# ---------------------------------------------------------------------------
try:
    import pdfplumber
    PDFPLUMBER_OK = True
except ImportError:
    PDFPLUMBER_OK = False

try:
    import openpyxl
    OPENPYXL_OK = True
except ImportError:
    OPENPYXL_OK = False

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


# ---------------------------------------------------------------------------
# Worker helpers (run in thread-pool workers)
# ---------------------------------------------------------------------------

def _extract_page(pdf_path: str, page_index: int):
    """
    Extract tables and plain text from a single PDF page.

    Returns a dict:
        {
          "page_num": int,          # 1-based page number
          "tables": list[list[list]], # list of tables, each table is rows of cells
          "text":   str,            # fallback plain text for the page
        }
    """
    with pdfplumber.open(pdf_path) as pdf:
        page = pdf.pages[page_index]
        tables = page.extract_tables()
        text = page.extract_text() or ""
    return {
        "page_num": page_index + 1,
        "tables": tables or [],
        "text": text,
    }


# ---------------------------------------------------------------------------
# Writers
# ---------------------------------------------------------------------------

def _write_excel(results: list, output_path: str) -> None:
    """Write extracted data to an Excel workbook."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Extraction"
    row_cursor = 1

    for result in results:
        page_num = result["page_num"]
        ws.cell(row=row_cursor, column=1, value=f"Page {page_num}")
        row_cursor += 1

        if result["tables"]:
            for table_idx, table in enumerate(result["tables"], start=1):
                if len(result["tables"]) > 1:
                    ws.cell(row=row_cursor, column=1, value=f"Table {table_idx}")
                    row_cursor += 1
                for row in table:
                    for col_idx, cell in enumerate(row, start=1):
                        ws.cell(row=row_cursor, column=col_idx,
                                value=cell if cell is not None else "")
                    row_cursor += 1
                row_cursor += 1  # blank row between tables
        else:
            # No tables found – write plain text instead
            for line in result["text"].splitlines():
                ws.cell(row=row_cursor, column=1, value=line)
                row_cursor += 1
        row_cursor += 1  # blank row between pages

    wb.save(output_path)


def _write_word(results: list, output_path: str) -> None:
    """Write extracted data to a Word document."""
    doc = DocxDocument()
    doc.add_heading("PDF Extraction Results", level=0)

    for result in results:
        page_num = result["page_num"]
        doc.add_heading(f"Page {page_num}", level=1)

        if result["tables"]:
            for table_idx, table in enumerate(result["tables"], start=1):
                if len(result["tables"]) > 1:
                    doc.add_paragraph(f"Table {table_idx}:", style="Intense Quote")

                if not table:
                    continue

                # Determine max columns
                num_cols = max(len(row) for row in table)
                word_table = doc.add_table(rows=len(table), cols=num_cols)
                word_table.style = "Table Grid"

                for r_idx, row in enumerate(table):
                    for c_idx, cell in enumerate(row):
                        word_table.cell(r_idx, c_idx).text = (
                            str(cell) if cell is not None else ""
                        )

                doc.add_paragraph()  # spacer
        else:
            doc.add_paragraph(result["text"] or "(No content found on this page)")

    doc.save(output_path)


def _write_txt(results: list, output_path: str) -> None:
    """Write extracted data to a plain-text file."""
    lines = []
    for result in results:
        page_num = result["page_num"]
        lines.append(f"{'=' * 60}")
        lines.append(f"  Page {page_num}")
        lines.append(f"{'=' * 60}")

        if result["tables"]:
            for table_idx, table in enumerate(result["tables"], start=1):
                if len(result["tables"]) > 1:
                    lines.append(f"\n  Table {table_idx}:")
                for row in table:
                    lines.append(
                        "\t".join(str(c) if c is not None else "" for c in row)
                    )
                lines.append("")
        else:
            lines.append(result["text"] or "(No content found on this page)")

        lines.append("")

    with open(output_path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))


# ---------------------------------------------------------------------------
# Main Application
# ---------------------------------------------------------------------------

class PDFExtractorApp:
    """Main GUI application for PDF table extraction."""

    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.root.title("PDF Table Extractor")
        self.root.resizable(False, False)

        # Try to set a reasonable minimum size
        self.root.minsize(560, 480)

        self._pdf_path: str | None = None
        self._total_pages: int = 0
        self._extraction_thread: threading.Thread | None = None
        self._progress_queue: queue.Queue = queue.Queue()

        self._build_ui()
        self._check_dependencies()

    # ------------------------------------------------------------------
    # UI Construction
    # ------------------------------------------------------------------

    def _build_ui(self) -> None:
        pad = {"padx": 10, "pady": 6}

        # ── Title ──────────────────────────────────────────────────────
        title_lbl = tk.Label(
            self.root,
            text="PDF Table Extractor",
            font=("Helvetica", 16, "bold"),
        )
        title_lbl.pack(**pad)

        # ── File Selection ─────────────────────────────────────────────
        file_frame = tk.LabelFrame(self.root, text="1. Select PDF File", padx=8, pady=6)
        file_frame.pack(fill="x", padx=12, pady=4)

        self._file_entry = tk.Entry(file_frame, width=48, state="readonly")
        self._file_entry.pack(side="left", padx=(0, 6))

        browse_btn = tk.Button(
            file_frame, text="Browse…", command=self._browse_file, width=10
        )
        browse_btn.pack(side="left")

        self._page_info_lbl = tk.Label(file_frame, text="", fg="grey")
        self._page_info_lbl.pack(side="left", padx=8)

        # ── Page Range ─────────────────────────────────────────────────
        range_frame = tk.LabelFrame(
            self.root, text="2. Page Range to Extract", padx=8, pady=6
        )
        range_frame.pack(fill="x", padx=12, pady=4)

        tk.Label(range_frame, text="From page:").grid(
            row=0, column=0, sticky="e", padx=4
        )
        self._from_var = tk.StringVar(value="1")
        from_spin = tk.Spinbox(
            range_frame,
            textvariable=self._from_var,
            from_=1,
            to=9999,
            width=6,
        )
        from_spin.grid(row=0, column=1, sticky="w", padx=4)

        tk.Label(range_frame, text="To page:").grid(
            row=0, column=2, sticky="e", padx=4
        )
        self._to_var = tk.StringVar(value="1")
        to_spin = tk.Spinbox(
            range_frame,
            textvariable=self._to_var,
            from_=1,
            to=9999,
            width=6,
        )
        to_spin.grid(row=0, column=3, sticky="w", padx=4)

        # ── Output Format ──────────────────────────────────────────────
        fmt_frame = tk.LabelFrame(
            self.root, text="3. Output Format", padx=8, pady=6
        )
        fmt_frame.pack(fill="x", padx=12, pady=4)

        self._format_var = tk.StringVar(value="excel")
        formats = [("Excel (.xlsx)", "excel"), ("Word (.docx)", "word"), ("Text (.txt)", "txt")]
        for col, (label, value) in enumerate(formats):
            rb = tk.Radiobutton(
                fmt_frame,
                text=label,
                variable=self._format_var,
                value=value,
            )
            rb.grid(row=0, column=col, padx=16, sticky="w")

        # ── Extract Button ─────────────────────────────────────────────
        self._extract_btn = tk.Button(
            self.root,
            text="⬇  Extract",
            font=("Helvetica", 11, "bold"),
            bg="#2d6cdf",
            fg="white",
            activebackground="#1a4faa",
            activeforeground="white",
            width=18,
            height=2,
            command=self._start_extraction,
        )
        self._extract_btn.pack(pady=8)

        # ── Progress ───────────────────────────────────────────────────
        prog_frame = tk.LabelFrame(self.root, text="Progress", padx=8, pady=6)
        prog_frame.pack(fill="x", padx=12, pady=4)

        self._progress_var = tk.DoubleVar(value=0.0)
        self._progress_bar = ttk.Progressbar(
            prog_frame,
            variable=self._progress_var,
            maximum=100,
            length=500,
            mode="determinate",
        )
        self._progress_bar.pack(pady=4)

        self._status_lbl = tk.Label(prog_frame, text="Ready.", anchor="w")
        self._status_lbl.pack(fill="x")

    # ------------------------------------------------------------------
    # Dependency Check
    # ------------------------------------------------------------------

    def _check_dependencies(self) -> None:
        missing = []
        if not PDFPLUMBER_OK:
            missing.append("pdfplumber")
        if not OPENPYXL_OK:
            missing.append("openpyxl")
        if not DOCX_OK:
            missing.append("python-docx")

        if missing:
            messagebox.showerror(
                "Missing Dependencies",
                "The following Python packages are required but not installed:\n\n"
                + "\n".join(f"  • {pkg}" for pkg in missing)
                + "\n\nRun:  pip install " + " ".join(missing),
            )

    # ------------------------------------------------------------------
    # File Browsing
    # ------------------------------------------------------------------

    def _browse_file(self) -> None:
        path = filedialog.askopenfilename(
            title="Select a PDF file",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
        )
        if not path:
            return

        self._pdf_path = path
        self._file_entry.config(state="normal")
        self._file_entry.delete(0, tk.END)
        self._file_entry.insert(0, path)
        self._file_entry.config(state="readonly")

        # Count pages and update UI
        try:
            with pdfplumber.open(path) as pdf:
                self._total_pages = len(pdf.pages)
            self._page_info_lbl.config(
                text=f"({self._total_pages} pages)", fg="green"
            )
            self._from_var.set("1")
            self._to_var.set(str(self._total_pages))
            self._status_lbl.config(text=f"Loaded: {os.path.basename(path)}")
        except Exception as exc:
            messagebox.showerror("Error", f"Could not open PDF:\n{exc}")
            self._pdf_path = None
            self._page_info_lbl.config(text="(error reading file)", fg="red")

    # ------------------------------------------------------------------
    # Extraction
    # ------------------------------------------------------------------

    def _start_extraction(self) -> None:
        if not PDFPLUMBER_OK:
            messagebox.showerror("Error", "pdfplumber is not installed.")
            return
        if not self._pdf_path:
            messagebox.showwarning("No File", "Please select a PDF file first.")
            return

        # Validate page range
        try:
            from_page = int(self._from_var.get())
            to_page = int(self._to_var.get())
        except ValueError:
            messagebox.showerror("Invalid Input", "Page numbers must be integers.")
            return

        if from_page < 1 or to_page < 1:
            messagebox.showerror("Invalid Input", "Page numbers must be ≥ 1.")
            return
        if from_page > to_page:
            messagebox.showerror(
                "Invalid Input", "'From' page must be ≤ 'To' page."
            )
            return
        if self._total_pages and to_page > self._total_pages:
            messagebox.showerror(
                "Invalid Input",
                f"'To' page ({to_page}) exceeds total pages ({self._total_pages}).",
            )
            return

        fmt = self._format_var.get()
        if fmt == "excel" and not OPENPYXL_OK:
            messagebox.showerror("Error", "openpyxl is not installed.")
            return
        if fmt == "word" and not DOCX_OK:
            messagebox.showerror("Error", "python-docx is not installed.")
            return

        # Disable button during extraction
        self._extract_btn.config(state="disabled")
        self._progress_var.set(0)
        self._status_lbl.config(text="Starting extraction…")

        # Run in background thread
        self._extraction_thread = threading.Thread(
            target=self._run_extraction,
            args=(self._pdf_path, from_page, to_page, fmt),
            daemon=True,
        )
        self._extraction_thread.start()
        self._poll_progress()

    def _run_extraction(
        self, pdf_path: str, from_page: int, to_page: int, fmt: str
    ) -> None:
        """Background worker: extract pages in parallel, then write output."""
        page_indices = list(range(from_page - 1, to_page))  # 0-based
        total = len(page_indices)
        results_map: dict = {}

        self._progress_queue.put(("status", f"Extracting pages {from_page}–{to_page}…"))

        # Parallel page extraction
        max_workers = min(4, total) if total > 1 else 1
        completed = 0

        try:
            with concurrent.futures.ThreadPoolExecutor(
                max_workers=max_workers
            ) as executor:
                future_to_idx = {
                    executor.submit(_extract_page, pdf_path, idx): idx
                    for idx in page_indices
                }

                for future in concurrent.futures.as_completed(future_to_idx):
                    result = future.result()
                    results_map[result["page_num"]] = result
                    completed += 1
                    pct = completed / total * 80  # first 80 % = extraction phase
                    self._progress_queue.put(
                        (
                            "progress",
                            pct,
                            f"Extracted page {result['page_num']} "
                            f"({completed}/{total} pages done)",
                        )
                    )
        except Exception as exc:
            self._progress_queue.put(("error", f"Extraction failed:\n{exc}"))
            return

        # Sort results by page number
        results = [results_map[pn] for pn in sorted(results_map.keys())]

        # Determine output path (same directory as PDF)
        self._progress_queue.put(("status", "Writing output file…"))
        base_dir = os.path.dirname(pdf_path)
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        range_suffix = f"_p{from_page}-{to_page}"

        ext_map = {"excel": ".xlsx", "word": ".docx", "txt": ".txt"}
        out_path = os.path.join(
            base_dir, base_name + range_suffix + ext_map[fmt]
        )

        try:
            if fmt == "excel":
                _write_excel(results, out_path)
            elif fmt == "word":
                _write_word(results, out_path)
            else:
                _write_txt(results, out_path)
        except Exception as exc:
            self._progress_queue.put(("error", f"Could not write output file:\n{exc}"))
            return

        self._progress_queue.put(("progress", 100, "Done!"))
        self._progress_queue.put(("done", out_path))

    # ------------------------------------------------------------------
    # Progress Polling
    # ------------------------------------------------------------------

    def _poll_progress(self) -> None:
        """Called periodically from the main thread to drain the progress queue."""
        try:
            while True:
                msg = self._progress_queue.get_nowait()
                tag = msg[0]

                if tag == "progress":
                    _, pct, status_text = msg
                    self._progress_var.set(pct)
                    self._status_lbl.config(text=status_text)

                elif tag == "status":
                    _, status_text = msg
                    self._status_lbl.config(text=status_text)

                elif tag == "error":
                    _, err_text = msg
                    self._status_lbl.config(text="Error – see dialog.")
                    self._extract_btn.config(state="normal")
                    messagebox.showerror("Extraction Error", err_text)
                    return  # stop polling

                elif tag == "done":
                    _, out_path = msg
                    self._extract_btn.config(state="normal")
                    messagebox.showinfo(
                        "Extraction Complete",
                        f"File saved to:\n{out_path}",
                    )
                    return  # stop polling

        except queue.Empty:
            pass

        # Schedule next poll in 100 ms if extraction is still running
        if self._extraction_thread and self._extraction_thread.is_alive():
            self.root.after(100, self._poll_progress)
        else:
            # Thread finished but queue might still have items – drain once more
            self._poll_progress()


# ---------------------------------------------------------------------------
# Entry Point
# ---------------------------------------------------------------------------

def main() -> None:
    root = tk.Tk()
    app = PDFExtractorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
