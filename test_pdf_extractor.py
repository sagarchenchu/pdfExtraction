"""
Tests for pdf_extractor business logic (runs without a display / tkinter).
"""

import os
import sys
import types
import unittest


# ---------------------------------------------------------------------------
# Stub tkinter so we can import pdf_extractor in headless CI environments
# ---------------------------------------------------------------------------

def _build_tkinter_stubs():
    class _Widget:
        def __init__(self, *a, **kw): pass
        def config(self, **kw): pass
        def pack(self, **kw): pass
        def grid(self, **kw): pass
        def delete(self, *a): pass
        def insert(self, *a): pass
        def after(self, *a): pass

    class _Var:
        def __init__(self, value=None): self._v = value
        def get(self): return self._v
        def set(self, v): self._v = v

    tk = types.ModuleType("tkinter")
    for name in ["Tk", "Frame", "LabelFrame", "Label", "Button", "Entry",
                 "Spinbox", "Radiobutton"]:
        setattr(tk, name, _Widget)
    tk.StringVar = _Var
    tk.DoubleVar = _Var

    ttk = types.ModuleType("tkinter.ttk")
    ttk.Progressbar = _Widget

    sys.modules.setdefault("tkinter", tk)
    sys.modules.setdefault("tkinter.ttk", ttk)
    sys.modules.setdefault("tkinter.filedialog", types.ModuleType("tkinter.filedialog"))
    sys.modules.setdefault("tkinter.messagebox", types.ModuleType("tkinter.messagebox"))


_build_tkinter_stubs()

REPO_ROOT = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, REPO_ROOT)

import pdf_extractor as app  # noqa: E402  (must come after stubs)
import openpyxl               # noqa: E402


# ---------------------------------------------------------------------------
# Sample data used across multiple tests
# ---------------------------------------------------------------------------

PAGE_WITH_TABLE = {
    "page_num": 2,
    "tables": [
        [
            ["Name", "Age", "City"],
            ["Alice", "30", "New York"],
            ["Bob", "25", "Los Angeles"],
        ]
    ],
    "text": "",
}

PAGE_TEXT_ONLY = {
    "page_num": 3,
    "tables": [],
    "text": "Some plain text on page 3.\nSecond line.",
}

PAGE_MULTI_TABLE = {
    "page_num": 4,
    "tables": [
        [["Col1", "Col2"], ["A", "B"]],
        [["X", "Y", "Z"], ["1", "2", "3"]],
    ],
    "text": "",
}


# ---------------------------------------------------------------------------
# Excel writer tests
# ---------------------------------------------------------------------------

class TestWriteExcel(unittest.TestCase):
    def _out(self, name):
        return f"/tmp/{name}"

    def test_creates_file(self):
        path = self._out("test_excel_creates.xlsx")
        app._write_excel([PAGE_WITH_TABLE], path)
        self.assertTrue(os.path.exists(path))

    def test_sheet_names_match_pages(self):
        path = self._out("test_excel_sheets.xlsx")
        app._write_excel([PAGE_WITH_TABLE, PAGE_TEXT_ONLY], path)
        wb = openpyxl.load_workbook(path)
        self.assertIn("Page 2", wb.sheetnames)
        self.assertIn("Page 3", wb.sheetnames)

    def test_table_data_written_correctly(self):
        path = self._out("test_excel_data.xlsx")
        app._write_excel([PAGE_WITH_TABLE], path)
        wb = openpyxl.load_workbook(path)
        ws = wb["Page 2"]
        self.assertEqual(ws.cell(1, 1).value, "Name")
        self.assertEqual(ws.cell(2, 1).value, "Alice")
        self.assertEqual(ws.cell(3, 2).value, "25")

    def test_text_fallback_when_no_table(self):
        path = self._out("test_excel_text_fallback.xlsx")
        app._write_excel([PAGE_TEXT_ONLY], path)
        wb = openpyxl.load_workbook(path)
        ws = wb["Page 3"]
        cell_values = [ws.cell(r, 1).value for r in range(1, ws.max_row + 1)]
        self.assertIn("Some plain text on page 3.", cell_values)

    def test_multiple_tables_per_page(self):
        path = self._out("test_excel_multi_table.xlsx")
        app._write_excel([PAGE_MULTI_TABLE], path)
        wb = openpyxl.load_workbook(path)
        ws = wb["Page 4"]
        # Both tables must appear – verify some cells exist
        all_values = set()
        for row in ws.iter_rows(values_only=True):
            all_values.update(v for v in row if v)
        self.assertIn("Col1", all_values)
        self.assertIn("X", all_values)


# ---------------------------------------------------------------------------
# Word writer tests
# ---------------------------------------------------------------------------

class TestWriteWord(unittest.TestCase):
    def test_creates_file(self):
        path = "/tmp/test_word_creates.docx"
        app._write_word([PAGE_WITH_TABLE], path)
        self.assertTrue(os.path.exists(path))

    def test_contains_heading(self):
        from docx import Document
        path = "/tmp/test_word_heading.docx"
        app._write_word([PAGE_WITH_TABLE, PAGE_TEXT_ONLY], path)
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        # Page heading should appear
        self.assertTrue(any("Page 2" in t for t in texts))

    def test_table_structure(self):
        from docx import Document
        path = "/tmp/test_word_table.docx"
        app._write_word([PAGE_WITH_TABLE], path)
        doc = Document(path)
        self.assertTrue(len(doc.tables) >= 1)
        first_table = doc.tables[0]
        self.assertEqual(first_table.cell(0, 0).text, "Name")
        self.assertEqual(first_table.cell(1, 0).text, "Alice")


# ---------------------------------------------------------------------------
# TXT writer tests
# ---------------------------------------------------------------------------

class TestWriteTxt(unittest.TestCase):
    def test_creates_file(self):
        path = "/tmp/test_txt_creates.txt"
        app._write_txt([PAGE_WITH_TABLE], path)
        self.assertTrue(os.path.exists(path))

    def test_page_header_present(self):
        path = "/tmp/test_txt_header.txt"
        app._write_txt([PAGE_WITH_TABLE, PAGE_TEXT_ONLY], path)
        with open(path) as fh:
            content = fh.read()
        self.assertIn("Page 2", content)
        self.assertIn("Page 3", content)

    def test_table_data_present(self):
        path = "/tmp/test_txt_table.txt"
        app._write_txt([PAGE_WITH_TABLE], path)
        with open(path) as fh:
            content = fh.read()
        self.assertIn("Alice", content)
        self.assertIn("New York", content)

    def test_plain_text_present(self):
        path = "/tmp/test_txt_plain.txt"
        app._write_txt([PAGE_TEXT_ONLY], path)
        with open(path) as fh:
            content = fh.read()
        self.assertIn("Some plain text on page 3.", content)

    def test_tab_separated_columns(self):
        path = "/tmp/test_txt_tabs.txt"
        app._write_txt([PAGE_WITH_TABLE], path)
        with open(path) as fh:
            content = fh.read()
        # Header row should be tab-separated
        self.assertIn("Name\tAge\tCity", content)


if __name__ == "__main__":
    unittest.main()
